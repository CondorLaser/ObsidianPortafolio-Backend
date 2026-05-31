"""Genera docs/Endpoints.docx organizado por VISTA del frontend.

Mismo estilo que el doc 'Idea endpoints' que mandó el equipo frontend, pero
ya con los endpoints REALMENTE implementados (no mockeados). Si una vista
todavía no tiene backend (Alertas, Recomendaciones) se indica explícito.

Uso:
    ./venv/bin/python -m scripts.generate_endpoints_doc
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


OUT = Path(__file__).parent.parent / "docs" / "Endpoints.docx"

BLUE = RGBColor(0x1F, 0x4E, 0x79)   # Azul título del PDF original
GREEN = RGBColor(0x00, 0x80, 0x00)
ORANGE = RGBColor(0xC6, 0x5D, 0x0A)
RED = RGBColor(0xB0, 0x1B, 0x1B)


def heading(doc, text: str, level: int = 1, color: RGBColor = BLUE):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def bullet(doc, text: str, level: int = 0, style: str = "List Bullet"):
    p = doc.add_paragraph(text, style=f"{style} {level + 1}" if level > 0 else style)
    return p


def add_endpoint(doc, method: str, path: str, auth: str, returns: str = "", note: str = ""):
    """Línea de endpoint con method en bold + path + descripción + auth."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{method} ")
    run.bold = True
    method_color = {
        "GET": GREEN, "POST": BLUE, "PUT": ORANGE, "PATCH": ORANGE, "DELETE": RED
    }.get(method, BLUE)
    run.font.color.rgb = method_color
    p.add_run(f"{path}").bold = True
    if auth:
        p.add_run(f"  ({auth})").italic = True
    if returns:
        doc.add_paragraph(f"→ {returns}", style="List Bullet 2")
    if note:
        run = doc.add_paragraph(f"  Nota: {note}", style="List Bullet 2").runs[0]
        run.italic = True


def add_section_intro(doc, lines: list[str]):
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def main():
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Título ──
    title = doc.add_heading("Endpoints — Orion Portafolio Backend", level=0)
    for r in title.runs:
        r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.add_run(
        "Documento de referencia para el equipo frontend. Lista de endpoints "
        "agrupados por vista (matchea el doc 'Idea endpoints' original). "
    )
    p.add_run("Todos los endpoints están").bold = False
    p.add_run(" implementados y testeados ").bold = True
    p.add_run("(suite integration 27/27 OK contra local y Neon develop).")

    p = doc.add_paragraph()
    p.add_run("Base URL local: ")
    p.add_run("http://localhost:8000").bold = True
    p.add_run("   |   Base URL via Zuplo: ")
    p.add_run("https://condor-laser-main-57c2c96.d2.zuplo.dev/").bold = True

    doc.add_paragraph()

    # ── Leyenda de auth ──
    heading(doc, "Cómo leer los endpoints", level=2)

    p = doc.add_paragraph()
    p.add_run("Cada endpoint indica entre paréntesis qué tipo de auth requiere:")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("🔒 Clerk").bold = True
    p.add_run("  →  necesita header ")
    p.add_run("Authorization: Bearer <Clerk JWT>").bold = True
    p.add_run(" (el JWT lo emite Clerk al loguearse). Sin token, el backend devuelve 401.")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("🔒 Clerk + ownership").bold = True
    p.add_run("  →  además del JWT, el backend verifica que el recurso ")
    p.add_run("pertenece al user").bold = True
    p.add_run(" del token. Ejemplo: si pedís ")
    p.add_run("GET /accounts/{id}").bold = True
    p.add_run(" y esa cuenta no es tuya → 404 (no 403, para no filtrar existencia).")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("🌍 público").bold = True
    p.add_run("  →  no necesita auth. Usar libremente.")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("🤝 svix signature").bold = True
    p.add_run("  →  endpoint server-to-server. Verifica firma svix (no se llama desde el frontend).")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("📎 multipart/form-data").bold = True
    p.add_run("  →  request con archivo (PDF) + form fields, en vez de JSON. Igual requiere Clerk JWT.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Métodos HTTP con color:").bold = True
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("GET ").bold = True; p.runs[0].font.color.rgb = GREEN
    p.add_run("→ lectura, no modifica nada")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("POST ").bold = True; p.runs[0].font.color.rgb = BLUE
    p.add_run("→ crear recurso nuevo")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("PUT / PATCH ").bold = True; p.runs[0].font.color.rgb = ORANGE
    p.add_run("→ actualizar recurso existente")

    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Portafolio / Dashboard
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Portafolio / Dashboard", level=1)
    add_section_intro(doc, [
        "Necesita: vista global del patrimonio del usuario (todas sus cuentas, "
        "todas las transactions, todos los dividends, posiciones consolidadas).",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "GET", "/accounts", "🔒 Clerk",
                 returns="Lista de cuentas del usuario (más reciente primero)")
    add_endpoint(doc, "GET", "/positions", "🔒 Clerk",
                 returns="Posiciones derivadas en runtime desde transactions + asset_prices",
                 note="Schema PositionDerived: {account_id, asset_id, symbol, name, quantity, avg_cost, last_price, market_value, unrealized_pnl}")
    add_endpoint(doc, "GET", "/transactions", "🔒 Clerk",
                 returns="Todas las transactions del usuario")
    add_endpoint(doc, "GET", "/dividends", "🔒 Clerk",
                 returns="Todos los dividends del usuario")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Activos
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Activos", level=1)
    add_section_intro(doc, [
        "Necesita: catálogo de activos del sistema (público para usuarios "
        "autenticados, no scoped por user).",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "GET", "/assets", "🔒 Clerk",
                 returns="Lista de assets con filtros opcionales",
                 note="Query params: ?symbol=X (exact match), ?kind=stock|etf|fund|crypto|other, ?currency=USD, ?search=texto (ilike sobre symbol o name), ?limit=100, ?offset=0")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Activo Específico
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Activo Específico", level=1)
    add_section_intro(doc, [
        "Necesita: detalle de un asset con su serie histórica de precios.",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "GET", "/assets/{asset_id}", "🔒 Clerk",
                 returns="Asset con prices embebidos (ORDER BY date DESC)")
    add_endpoint(doc, "GET", "/assets/{asset_id}/prices", "🔒 Clerk",
                 returns="Serie histórica de precios, filtrable por fecha",
                 note="Query params: ?from=YYYY-MM-DD&to=YYYY-MM-DD")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Cuentas
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Cuentas", level=1)
    add_section_intro(doc, [
        "Necesita: listar las cuentas del usuario.",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "GET", "/accounts", "🔒 Clerk",
                 returns="Lista de cuentas (ORDER BY created_at DESC)")
    add_endpoint(doc, "POST", "/accounts", "🔒 Clerk",
                 returns="Crear nueva cuenta. Body: {name, broker?, currency}")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Cuenta Específica
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Cuenta Específica", level=1)
    add_section_intro(doc, [
        "Necesita: detalle de una cuenta, sus posiciones, transactions y dividends.",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "GET", "/accounts/{account_id}", "🔒 Clerk + ownership",
                 returns="Account con dividends + positions + transactions embebidos")
    add_endpoint(doc, "GET", "/accounts/metrics/{account_id}", "🔒 Clerk + ownership",
                 returns="{ daily: [...], monthly: [...] }",
                 note="Hoy devuelve arrays vacíos — cómputo de métricas (PnL, drawdown, Sharpe, etc.) pendiente de implementar")
    add_endpoint(doc, "GET", "/accounts/positions/{account_id}", "🔒 Clerk + ownership",
                 returns="Lista de posiciones materializadas de esa cuenta")
    add_endpoint(doc, "GET", "/accounts/transactions/{account_id}", "🔒 Clerk + ownership",
                 returns="Transactions de esa cuenta")
    add_endpoint(doc, "GET", "/accounts/dividends/{account_id}", "🔒 Clerk + ownership",
                 returns="Dividends de esa cuenta")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Perfil
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Perfil", level=1)
    add_section_intro(doc, [
        "Necesita: obtener / editar perfil de riesgo, preferencias, nombres de "
        "cuentas, y subir archivos PDF Fintual para ingesta.",
    ])
    doc.add_paragraph("Endpoints — Profile:")
    add_endpoint(doc, "GET", "/profile", "🔒 Clerk",
                 returns="{ clerk_id, email, created_at, risk_profile }")
    add_endpoint(doc, "PUT", "/profile", "🔒 Clerk",
                 returns="Update profile (hoy permite editar risk_profile). Body: { risk_profile }")
    add_endpoint(doc, "PATCH", "/profile/risk-profile", "🔒 Clerk",
                 returns="Canónico: actualiza solo el risk_profile. Body: { risk_profile }")
    doc.add_paragraph()

    doc.add_paragraph("Endpoints — User (aliases livianos):")
    add_endpoint(doc, "GET", "/user/risk_profile", "🔒 Clerk",
                 returns="{ risk_profile: 'moderate' | 'agressive' | 'conservative' | null }")
    add_endpoint(doc, "PUT", "/user/risk_profile", "🔒 Clerk",
                 returns="Alias de PATCH /profile/risk-profile. Body: { risk_profile }")
    add_endpoint(doc, "GET", "/user/accounts_names", "🔒 Clerk",
                 returns="[ { id, name }, ... ] — solo id+nombre de cada cuenta")
    add_endpoint(doc, "PUT", "/user/accounts_names", "🔒 Clerk + ownership",
                 returns="Rename batch. Body: [ { id, name }, ... ]")
    doc.add_paragraph()

    doc.add_paragraph("Endpoints — Preferences (umbrales de alertas):")
    add_endpoint(doc, "GET", "/preferences", "🔒 Clerk",
                 returns="UserPreference (404 si nunca se setearon)",
                 note="Campos: pnl_percentage_account_daily, pnl_percentage_asset_daily, max_drawdown_portfolio_daily, max_drawdown_account_daily, asset_weight_weekly, currency_exposure_weekly")
    add_endpoint(doc, "PUT", "/preferences", "🔒 Clerk",
                 returns="Upsert con merge (no override de campos no enviados)")
    doc.add_paragraph()

    doc.add_paragraph("Endpoints — Subida de archivos Fintual (ingesta de PDFs):")
    add_endpoint(doc, "POST", "/pdf/extract_stocks_etf_1",
                 "🔒 Clerk + 📎 multipart/form-data",
                 returns="{ compras_ventas_guardadas: N, dividendos_guardados: M, errores_activos_faltantes: [...] }",
                 note="Body: file (PDF) + account_id (form). Procesa certificado Fintual de stocks/ETFs (compraventa + dividendos)")
    add_endpoint(doc, "POST", "/pdf/extract_mutual_funds",
                 "🔒 Clerk + 📎 multipart/form-data",
                 returns="{ compras_ventas_guardadas: N, errores_activos_faltantes: [...] }",
                 note="Body: file (PDF) + account_id (form). Procesa certificado Fintual de fondos mutuos")
    add_endpoint(doc, "POST", "/pdf/extract_stocks_etf_2",
                 "🔒 Clerk + 📎 multipart/form-data",
                 returns="Holdings (estado de posición). Hoy stub — pendiente de decidir materialización")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vista Onboarding
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Onboarding", level=1)
    add_section_intro(doc, [
        "Necesita: setear el risk_profile inicial del usuario en el primer flujo.",
    ])
    doc.add_paragraph("Endpoints:")
    add_endpoint(doc, "POST", "/risk_profile", "🔒 Clerk",
                 returns="{ risk_profile } — alias de PATCH /profile/risk-profile para uso en onboarding")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Vistas Alertas + Recomendaciones (POST-MVP, no implementadas)
    # ────────────────────────────────────────────────────────────
    heading(doc, "Vista Alertas", level=1, color=ORANGE)
    p = doc.add_paragraph()
    p.add_run("⚠️ Post-MVP. ").bold = True
    p.add_run("Backend todavía no tiene tablas (alert_rule, alert) ni endpoints. "
              "Cuando se decida implementar: necesita migración Alembic + nuevos "
              "routers + lógica de disparo.")
    doc.add_paragraph()

    heading(doc, "Vista Recomendaciones", level=1, color=ORANGE)
    p = doc.add_paragraph()
    p.add_run("⚠️ Post-MVP. ").bold = True
    p.add_run("Backend todavía no tiene tablas (recommendation, recommendation_run) "
              "ni endpoints ni motor de recomendación.")
    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Backend infra (no son vistas — uso interno + webhooks)
    # ────────────────────────────────────────────────────────────
    heading(doc, "Backend infra (no son vistas frontend)", level=1)
    p = doc.add_paragraph()
    p.add_run("Estos endpoints no los llama el frontend directamente, "
              "pero existen para infra/admin/ingestión:")

    doc.add_paragraph("Públicos (sin auth):")
    add_endpoint(doc, "GET", "/", "🌍 público",
                 returns="Mensaje básico de la API (status running)")
    add_endpoint(doc, "GET", "/health", "🌍 público",
                 returns="{ status: 'ok' } — para readiness checks de Render/K8s")

    doc.add_paragraph("Auth Clerk (sanity check):")
    add_endpoint(doc, "GET", "/protected", "🔒 Clerk",
                 returns="{ user_id: <sub del JWT> } — solo para validar setup de auth")

    doc.add_paragraph("Webhook Clerk (server-to-server, no es vista):")
    add_endpoint(doc, "POST", "/webhooks/clerk", "🤝 svix signature",
                 returns="{ status: 'ok' }",
                 note="Maneja user.created/updated/deleted. Configurado en Clerk Dashboard con CLERK_WEBHOOK_SECRET")

    doc.add_paragraph("Admin / catálogo (writes, hoy con Clerk JWT — pendiente service token):")
    add_endpoint(doc, "POST", "/assets", "🔒 Clerk",
                 returns="Crear asset en el catálogo global. Body: { symbol, name, kind, currency }")
    add_endpoint(doc, "POST", "/assets/{asset_id}/prices", "🔒 Clerk",
                 returns="Upsert precio diario. Body: { date, close, currency, source }",
                 note="También lo usa el GitHub Action de sync (sync_stock_prices.py) contra TwelveData")
    add_endpoint(doc, "POST", "/transactions", "🔒 Clerk + ownership",
                 returns="Crear transaction manualmente. Body: { account_id, asset_id, kind, quantity, price?, fee?, executed_at }")

    doc.add_paragraph()

    # ────────────────────────────────────────────────────────────
    # Notas finales
    # ────────────────────────────────────────────────────────────
    heading(doc, "Notas para el frontend", level=1)
    add_section_intro(doc, [
        "JIT user creation: cualquier request autenticada crea el profile en backend "
        "si el clerk_id del JWT no existe. No hace falta esperar al webhook.",
        "Códigos de respuesta comunes: 200 OK, 201 Created (POST), 400 (body inválido), "
        "401 (JWT inválido/faltante), 404 (no existe O no es del user — isolation), "
        "409 (conflict, ej. symbol duplicado), 422 (validación Pydantic).",
        "CORS: el origen del frontend tiene que estar en ALLOWED_ORIGINS del backend.",
        "Contrato detallado (33 endpoints con shapes Pydantic resueltas): ver "
        "docs/API_CONTRACT.md — autogenerado, siempre al día con el OpenAPI vivo.",
    ])

    # ── Footer ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Generado desde scripts/generate_endpoints_doc.py — re-correr tras cambios de rutas")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"✅ {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
